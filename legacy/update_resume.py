from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================================================================
# DATA DICTIONARY — fill in actual content here before running
# ==============================================================================

DATA = {
    "header": {
        "name_block": (
            "Farhan Solaiman\n"
            "Mezza Residences, Aurora Blvd. Cor. Araneta Ave. \u00b7 09271870882\n"
            "farhant.solaiman@gmail.com\n"
            "Portfolio: farhansolaiman.github.io/developer-portfolio/"
        ),
        "summary": (
            "Senior Full Stack Developer with 7+ years of experience building "
            "scalable applications, cloud infrastructure, and intelligent automation "
            "systems. Skilled in React, Node.js, TypeScript, AWS, and AI workflow "
            "automation. Proven track record of driving technical innovation, leading "
            "cross-functional teams, and delivering high-impact solutions that improve "
            "operational efficiency by 25\u201360%."
        ),
    },
    "experience": [
        {
            "dates": "October 2025 \u2013 Present",
            "title": "AI Automation Tech Lead, ",
            "company": "AI Automation Agency",
            "sub_entries": [
                {
                    "header": None,
                    "bullets": [
                        "Architected and deployed 15+ production AI-powered automation workflows using n8n, integrating OpenAI and Gemini APIs to reduce manual processing time by up to 60%.",
                        'Built "Citation Sniper," an end-to-end AI pipeline combining SerpAPI, Browserless.io, and GPT-4 for automated SEO citation discovery and competitor analysis.',
                        'Designed "ClickUp Agency OS," connecting ClickUp, Slack, and email workflows via n8n for streamlined agency operations and real-time team collaboration.',
                        "Led a team of 5 engineers using sprint-based delivery, establishing code review standards and building a reusable automation template library that accelerated delivery by 40%.",
                        "Implemented AI-driven data pipelines with headless browser automation and LLM-powered data extraction, establishing monitoring patterns achieving 99%+ workflow uptime.",
                        "Managed full project lifecycle from client discovery through production deployment, ensuring scalability and maintainability across all automation solutions.",
                    ],
                }
            ],
        },
        {
            "dates": "April 2023 \u2013 September 2025",
            "title": "Senior Full Stack Developer, ",
            "company": "AS White Global",
            "sub_entries": [
                {
                    "header": None,
                    "bullets": [
                        "Improved UI performance by 30% through optimized React and TypeScript implementations, enhancing user experience and application responsiveness.",
                        "Designed RESTful APIs that reduced data transfer errors by 25%, improving system reliability and data integrity across client platforms.",
                        "Partnered with product owners and QA to deliver features 2 weeks ahead of schedule by implementing TDD practices and establishing rigorous code review standards.",
                        "Mentored junior developers, improving onboarding efficiency and establishing best practices for full-stack development.",
                        "Supported AWS cloud migration initiatives, contributing to cost reduction and improved scalability for enterprise applications.",
                    ],
                }
            ],
        },
        {
            "dates": "July 2021 \u2013 April 2023",
            "title": "Full Stack Engineer, ",
            "company": "Accenture",
            "sub_entries": [
                {
                    "header": "Energy Provider Company \u2013 [October 2022 \u2013 April 2023]",
                    "bullets": [
                        "Developed and integrated custom backend APIs with frontend client applications, delivering seamless end-to-end functionality.",
                        "Authored Terraform Infrastructure-as-Code (IaC) scripts to provision and manage AWS services, streamlining deployment and improving scalability.",
                        "Modified and optimized CI/CD pipelines, reducing release cycle times and aligning deployments with evolving system requirements.",
                    ],
                },
                {
                    "header": "Energy Provider Company \u2013 [July 2022 \u2013 October 2022]",
                    "bullets": [
                        "Led a team of 3 full-stack developers, providing technical guidance and oversight across both frontend and backend tasks.",
                        "Built analytics and data visualization solutions using Power BI, tailored to stakeholder reporting needs.",
                        "Automated the full lifecycle of data retrieval and processing using Excel Office Scripts and Power Automate, significantly reducing manual work.",
                    ],
                },
                {
                    "header": "Energy Provider Company \u2013 [August 2021 \u2013 July 2022]",
                    "bullets": [
                        "Designed and implemented multiple custom APIs to support new features and business requirements.",
                        "Developed responsive UI components from PDF/Figma design specifications, delivering pixel-perfect interfaces.",
                        "Utilized Jest to write and maintain unit tests, ensuring code quality and functional integrity.",
                        "Achieved AWS Certified Cloud Practitioner certification.",
                    ],
                },
                {
                    "header": "Refining and Manufacturing Company \u2013 [April 2021 \u2013 July 2021]",
                    "bullets": [
                        "Developed serverless functions using TypeScript and Node.js, leveraging AWS services (Lambda, SQS, S3, CloudFormation, CloudWatch, IAM).",
                        "Created a custom database checker application to identify potential database issues early, accelerating team progress.",
                        "Provided mentorship and technical support to junior developers through code reviews and collaborative debugging sessions.",
                    ],
                },
            ],
        },
        {
            "dates": "October 2018 \u2013 March 2021",
            "title": "Science Research Specialist I (Web Developer), ",
            "company": "DOST-ASTI",
            "sub_entries": [
                {
                    "header": None,
                    "bullets": [
                        "Modernized legacy systems by migrating critical applications to web-based platforms, reducing manual processing time by 35%.",
                        "Developed secure internal tools using JavaScript, Node.js, and SQL, improving data accessibility and reporting accuracy for stakeholders.",
                        "Provided technical support and system enhancements that minimized downtime and increased service reliability for department users.",
                        "Partnered with cross-functional teams to align technical solutions with government compliance standards and operational needs.",
                    ],
                }
            ],
        },
    ],
    "skills": {
        "col1": [
            "Full Stack Development",
            "AI Automation & Workflows",
            "API Design & Integration",
            "Cloud Architecture",
            "Code Quality & Testing",
            "Performance Optimization",
            "Team Leadership & Mentoring",
        ],
        "col2": [
            "Serverless Architecture",
            "Data Pipeline Engineering",
            "Headless Browser Automation",
            "Database Design",
            "Infrastructure as Code",
            "Agile/Sprint Delivery",
            "Technical Documentation",
        ],
    },
    "tools": {
        "col1": [
            "React",
            "Node.js",
            "TypeScript",
            "JavaScript",
            "AWS (Lambda, SQS, S3, CloudWatch, IAM)",
            "n8n",
            "OpenAI",
            "Gemini API",
            "SerpAPI",
            "Browserless.io",
        ],
        "col2": [
            "Express",
            "REST APIs",
            "Terraform",
            "MySQL",
            "MongoDB",
            "Redis",
            "Jest",
            "CircleCI",
            "GitHub",
            "GitLab",
        ],
    },
    "education": [
        {
            "date": "May 2018",
            "text": "BS Electronics Engineering, University of the East",
        },
        {
            "date": "September 2018",
            "text": "Full-Stack Web Development Bootcamp, Zuitt Coding Bootcamp",
            "note": "Top Performer of the class",
        },
    ],
    "certificates": [
        "AWS Certified Cloud Practitioner",
        "Electronics Technician",
        "Electronics Engineer",
    ],
}

# ==============================================================================
# HELPER FUNCTIONS
# ==============================================================================

def clear_cell(cell):
    """
    Remove all paragraphs from a cell except the first, then clear that one.
    python-docx requires at least one paragraph per cell, so we reuse the first.
    """
    paragraphs = cell.paragraphs
    # Remove all paragraphs after the first by clearing their XML elements
    for para in paragraphs[1:]:
        p_element = para._element
        p_element.getparent().remove(p_element)
    # Clear the first paragraph: remove all runs and reset style
    first_para = cell.paragraphs[0]
    for run in first_para.runs:
        run_element = run._element
        run_element.getparent().remove(run_element)
    first_para.style = cell._element.getroottree().getroot().body.getparent() \
        if False else first_para.style  # no-op placeholder
    # Reset style to Normal and clear any leftover text
    first_para.style = "Normal"
    first_para.clear()


def _apply_first_or_add(cell, is_first_call_tracker):
    """
    Internal helper: returns the first paragraph (modifying in place) on the
    very first call, then uses cell.add_paragraph() for subsequent calls.

    is_first_call_tracker is a one-element list [bool] used as a mutable flag.
    """
    if is_first_call_tracker[0]:
        is_first_call_tracker[0] = False
        return cell.paragraphs[0]
    else:
        return cell.add_paragraph()


def add_date_paragraph(cell, text, _tracker=None):
    """Add a paragraph with Heading 3 style for date ranges."""
    if _tracker is None:
        para = cell.add_paragraph()
    else:
        para = _apply_first_or_add(cell, _tracker)
    para.style = "Heading 3"
    para.clear()
    run = para.add_run(text)
    return para


def add_title_paragraph(cell, title_text, company_text, _tracker=None):
    """
    Add a paragraph with Heading 2 style.
    title_text uses default style runs; company_text uses a separate run with
    bold=False and color RGBColor(0x59, 0x59, 0x59).
    """
    if _tracker is None:
        para = cell.add_paragraph()
    else:
        para = _apply_first_or_add(cell, _tracker)
    para.style = "Heading 2"
    para.clear()
    # Title run — inherits Heading 2 styling
    title_run = para.add_run(title_text)
    # Company run — override bold and color
    company_run = para.add_run(company_text)
    company_run.bold = False
    company_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return para


def add_subheader_paragraph(cell, text, _tracker=None):
    """
    Add a Normal paragraph for sub-period headers.
    Run: bold=True, italic=False, size=Pt(11), color=RGBColor(0x59, 0x59, 0x59).
    """
    if _tracker is None:
        para = cell.add_paragraph()
    else:
        para = _apply_first_or_add(cell, _tracker)
    para.style = "Normal"
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.italic = False
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return para


def add_bullet_paragraph(cell, text, _tracker=None):
    """
    Add a Normal paragraph with LEFT alignment for bullet descriptions.
    Run: bold=False, italic=False, size=Pt(11), color=RGBColor(0x59, 0x59, 0x59).
    """
    if _tracker is None:
        para = cell.add_paragraph()
    else:
        para = _apply_first_or_add(cell, _tracker)
    para.style = "Normal"
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = False
    run.italic = False
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    return para


def add_empty_paragraph(cell, _tracker=None):
    """Add an empty Normal paragraph as a separator."""
    if _tracker is None:
        para = cell.add_paragraph()
    else:
        para = _apply_first_or_add(cell, _tracker)
    para.style = "Normal"
    para.clear()
    return para


def add_simple_text(cell, text, _tracker=None):
    """Add a Normal paragraph with plain text."""
    if _tracker is None:
        para = cell.add_paragraph()
    else:
        para = _apply_first_or_add(cell, _tracker)
    para.style = "Normal"
    para.clear()
    para.add_run(text)
    return para


# ==============================================================================
# SECTION BUILDERS
# ==============================================================================

def build_header(table, data):
    """Rebuild Table 0: header name/contact block and summary."""
    print("  Building header (Table 0)...")

    # Row 0, col 0 — name and contact block
    cell = table.rows[0].cells[0]
    clear_cell(cell)
    tracker = [True]
    first_para = _apply_first_or_add(cell, tracker)
    first_para.style = "Normal"
    first_para.clear()
    first_para.add_run(data["name_block"])

    # Row 1, col 0 — summary
    cell = table.rows[1].cells[0]
    clear_cell(cell)
    tracker = [True]
    add_simple_text(cell, data["summary"], _tracker=tracker)

    print("  Header built.")


def build_experience(table, entries):
    """Rebuild Table 1, row 0: all experience entries."""
    print("  Building experience (Table 1)...")

    cell = table.rows[0].cells[0]
    clear_cell(cell)
    tracker = [True]  # mutable flag — first paragraph must be reused in place

    for i, entry in enumerate(entries):
        print(f"    Adding experience entry {i + 1}: {entry.get('company', '')}")

        # Date line
        add_date_paragraph(cell, entry["dates"], _tracker=tracker)

        # Title + company line
        add_title_paragraph(
            cell,
            entry["title"],
            entry["company"],
            _tracker=tracker
        )

        # Sub-entries (sub-period headers + bullets)
        for sub in entry.get("sub_entries", []):
            if sub.get("header"):
                add_subheader_paragraph(cell, sub["header"], _tracker=tracker)
            for bullet in sub.get("bullets", []):
                add_bullet_paragraph(cell, bullet, _tracker=tracker)

        # Empty separator between entries (skip after last)
        if i < len(entries) - 1:
            add_empty_paragraph(cell, _tracker=tracker)

    print("  Experience built.")


def build_skills(table, skills_data):
    """Rebuild Table 2: skills (2 cols)."""
    print("  Building skills (Table 2)...")

    row = table.rows[0]
    for col_key, cell in zip(["col1", "col2"], row.cells):
        clear_cell(cell)
        tracker = [True]
        for item in skills_data.get(col_key, []):
            add_simple_text(cell, item, _tracker=tracker)

    print("  Skills built.")


def build_tools(table, tools_data):
    """Rebuild Table 3: languages and tools (2 cols)."""
    print("  Building tools (Table 3)...")

    row = table.rows[0]
    for col_key, cell in zip(["col1", "col2"], row.cells):
        clear_cell(cell)
        tracker = [True]
        for item in tools_data.get(col_key, []):
            add_simple_text(cell, item, _tracker=tracker)

    print("  Tools built.")


def build_education(table, education_entries):
    """Rebuild Table 4: education entries (rows 0 and 1)."""
    print("  Building education (Table 4)...")

    for row_idx, entry in enumerate(education_entries):
        if row_idx >= len(table.rows):
            print(f"    Warning: more education entries than rows — skipping entry {row_idx + 1}")
            break
        cell = table.rows[row_idx].cells[0]
        clear_cell(cell)
        tracker = [True]

        add_date_paragraph(cell, entry["date"], _tracker=tracker)
        add_title_paragraph(cell, entry["text"], "", _tracker=tracker)

        if entry.get("note"):
            add_bullet_paragraph(cell, entry["note"], _tracker=tracker)

    print("  Education built.")


def build_certificates(table, certificates):
    """Rebuild Table 5: certificates."""
    print("  Building certificates (Table 5)...")

    cell = table.rows[0].cells[0]
    clear_cell(cell)
    tracker = [True]

    for cert in certificates:
        add_simple_text(cell, cert, _tracker=tracker)

    print("  Certificates built.")


# ==============================================================================
# MAIN
# ==============================================================================

def main():
    filename = "farhan_solaiman_resume_2026.docx"
    print(f"Loading document: {filename}")
    doc = Document(filename)

    tables = doc.tables
    print(f"Found {len(tables)} tables in document.")

    if len(tables) < 6:
        raise ValueError(
            f"Expected at least 6 tables, found {len(tables)}. "
            "Check that the input file is the correct document."
        )

    print("\nRebuilding document sections...")

    build_header(tables[0], DATA["header"])
    build_experience(tables[1], DATA["experience"])
    build_skills(tables[2], DATA["skills"])
    build_tools(tables[3], DATA["tools"])
    build_education(tables[4], DATA["education"])
    build_certificates(tables[5], DATA["certificates"])

    print(f"\nSaving modified document to: {filename}")
    doc.save(filename)
    print("Done.")


if __name__ == "__main__":
    main()
